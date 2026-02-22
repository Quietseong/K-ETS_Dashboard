import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO

# 기존 호환성 유지: get_data_context는 data_loader에서 제공
from data_loader import get_data_context  # noqa: F401


# --- 보고서 파일 생성 유틸리티 ---

def create_docx(title: str, content: str) -> BytesIO:
    """
    보고서 제목과 내용을 받아 DOCX 파일 객체를 생성합니다.

    Args:
        title (str): 보고서의 제목.
        content (str): 보고서의 전체 본문 내용.

    Returns:
        BytesIO: 메모리 상에 생성된 DOCX 파일 객체.
    """
    document = Document()
    document.add_heading(title, level=1)

    # 내용을 문단별로 나누어 추가
    for paragraph in content.split('\n'):
        if paragraph.strip():
            document.add_paragraph(paragraph)

    # 파일 객체를 메모리에 저장
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_pdf(title: str, content: str) -> BytesIO:
    """
    보고서 제목과 내용을 받아 PDF 파일 객체를 생성합니다.
    한글을 지원하기 위해 'NanumGothic' 폰트를 사용합니다.

    Args:
        title (str): 보고서의 제목.
        content (str): 보고서의 전체 본문 내용.

    Returns:
        BytesIO: 메모리 상에 생성된 PDF 파일 객체.
    """
    # 폰트 등록 (파일 경로를 확인해야 할 수 있음, 임시로 상대 경로 사용)
    # 실제 배포 환경에서는 폰트 파일의 경로를 정확히 지정해야 합니다.
    try:
        import platform
        system = platform.system()

        if system == "Windows":
            font_paths = ["c:/Windows/Fonts/malgun.ttf", "c:/Windows/Fonts/gulim.ttc"]
        elif system == "Darwin":  # macOS
            font_paths = ["/System/Library/Fonts/AppleGothic.ttf"]
        else:  # Linux
            font_paths = ["/usr/share/fonts/truetype/nanum/NanumGothic.ttf"]

        font_name = 'Helvetica'  # fallback
        for font_path in font_paths:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('KoreanFont', font_path))
                font_name = 'KoreanFont'
                break
    except Exception:
        # 폰트 로드 실패 시 기본 폰트로 대체 (한글 깨질 수 있음)
        print("경고: '맑은 고딕' 폰트를 찾을 수 없습니다. PDF에서 한글이 깨질 수 있습니다.")
        font_name = 'Helvetica'

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    # 제목
    p.setFont(font_name, 16)
    p.drawString(50, height - 50, title)

    # 내용
    p.setFont(font_name, 10)
    text = p.beginText(50, height - 100)
    text.setLeading(15)

    for line in content.split('\n'):
        text.textLine(line)

    p.drawText(text)
    p.showPage()
    p.save()

    buffer.seek(0)
    return buffer
